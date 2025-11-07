"""
Word 문서 파서
"""
from docx import Document
from datetime import datetime
from typing import List

from app.parsers.base import DocumentParser, ParsedDocument, DocumentMetadata, ContentChunk


class WordParser(DocumentParser):
    """Word 문서 파서"""
    
    def parse(self, file_path: str) -> ParsedDocument:
        """Word 문서 파싱"""
        doc = Document(file_path)
        
        full_text = ""
        chunks = []
        structure = {
            "paragraphs": [],
            "tables": [],
            "sections": []
        }
        
        current_section = ""
        
        # 단락 추출
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 제목 스타일 감지
            if para.style.name.startswith('Heading'):
                current_section = text
                structure["sections"].append({
                    "title": text,
                    "level": para.style.name.replace('Heading ', '')
                })
            else:
                full_text += text + "\n"
            
            structure["paragraphs"].append({
                "text": text,
                "style": para.style.name,
                "section": current_section
            })
        
        # 표 추출
        for table_idx, table in enumerate(doc.tables):
            table_text = self._table_to_text(table)
            full_text += table_text + "\n"
            structure["tables"].append({
                "table_index": table_idx,
                "rows": len(table.rows)
            })
        
        # 메타데이터 추출
        metadata = self.extract_metadata(file_path)
        metadata.word_count = len(full_text.split())
        
        # 청크 분할
        chunks = self.chunk_document(full_text)
        
        # 섹션 정보 추가
        for chunk in chunks:
            chunk.section_title = current_section
        
        return ParsedDocument(
            filename=file_path.split("/")[-1],
            file_type="docx",
            metadata=metadata,
            chunks=chunks,
            full_text=full_text,
            structure=structure
        )
    
    def extract_metadata(self, file_path: str) -> DocumentMetadata:
        """Word 메타데이터 추출"""
        metadata = DocumentMetadata()
        
        try:
            doc = Document(file_path)
            core_props = doc.core_properties
            
            metadata.title = core_props.title or ""
            metadata.author = core_props.author or ""
            
            if core_props.created:
                metadata.created_date = core_props.created
            if core_props.modified:
                metadata.modified_date = core_props.modified
        except Exception as e:
            print(f"메타데이터 추출 오류: {e}")
        
        return metadata
    
    def _table_to_text(self, table) -> str:
        """표를 텍스트로 변환"""
        text_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            text_row = " | ".join(cells)
            text_rows.append(text_row)
        return "\n".join(text_rows)

